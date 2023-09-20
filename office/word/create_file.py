import docx
from docx.shared import Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# создаем объект документа
doc = docx.Document()

# добавление заголовка Текст и Уровень
doc.add_heading('Заголовок 0', 0)  # первый уровень
doc.add_heading('Заголовок 1', 1)  # второй уровень
h3 = doc.add_heading('Заголовок 2', 2)  # третий уровень
h3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# добавление параграфа Текст параграфа
doc.add_paragraph('Здравствуй, мир!')
abzac1 = doc.add_paragraph(text='Это второй абзац.')

# добавляем текст во второй параграф
r1 = abzac1.add_run('Этот текст был добавлен во второй абзац.', )
# добавить к этому тексту стиль Курсив (italic, bold, strike, underline)
r1.italic = True

abzac2 = doc.add_paragraph()
abzac3 = doc.add_paragraph('Текст для четвертого параграфа', 'Body Text 3')
# стиль для параграфа
#
abzac2.style = 'List 3'
r2 = abzac2.add_run('Этот текст был добавлен в третий абзац.')
# стиль для run
# Emphasis, Strong, Book Title, Default Paragraph Font, Intense Emphasis, Subtle Emphasis
# Intense Reference, Subtle Reference
r2.style = 'Book Title'
r3 = abzac2.add_run('Этот текст был добавлен в третий абзац.', 'Intense Reference')

# добавить разрыв страниц
doc.paragraphs[0].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
doc.add_page_break()

# добавить каритнку
doc.add_picture('1.jpg', width=Cm(10))

# добавляем таблицу 3x3
table = doc.add_table(rows=3, cols=3)
# применяем стиль для таблицы
table.style = 'Table Grid'

# заполняем таблицу данными
for row in range(3):
    for col in range(3):
        # получаем ячейку таблицы
        cell = table.cell(row, col)
        # записываем в ячейку данные
        cell.text = str(row + 1) + str(col + 1)


# маркированный список
doc.add_paragraph('This is the first item.', style='List Bullet')
doc.add_paragraph('This is the second item.', style='List Bullet')

# нумерованный список
doc.add_paragraph('This is the first item.', style='List Number')
doc.add_paragraph('This is the second item.', style='List Number')
doc.add_paragraph('This is the third item.', style='List Number')

# сохраняем файл Название файла
doc.save('filename.docx')
